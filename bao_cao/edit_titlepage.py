# -*- coding: utf-8 -*-
"""Them dong 'Nganh hoc: ET-LUH' va can giua bang thong tin trang bia,
giu nguyen toan bo noi dung khac nguoi dung da tu sua."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = r"C:\Users\HI\Documents\rtk tt\BAO_CAO_THUC_TAP.docx"
FONT = "Times New Roman"

doc = Document(PATH)
table = doc.tables[0]


def set_run_font(run, size=13, bold=False, name=FONT):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=13, bold=bold)


# Them dong "Nganh hoc: ET-LUH" vao cuoi bang
new_row = table.add_row()
set_cell(new_row.cells[0], "Ngành học:", bold=True)
set_cell(new_row.cells[1], "ET-LUH")

# Can giua ca bang tren trang (nhu cac dong tieu de khac cua trang bia)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Dat lai chieu rong cot cho dep (ap dung cho ca dong moi them)
for row in table.rows:
    row.cells[0].width = Cm(4.5)
    row.cells[1].width = Cm(9)

doc.save(PATH)
print("Da luu:", PATH)
